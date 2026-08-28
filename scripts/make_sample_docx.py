"""Build a sample advert + email sequence .docx for testing the parser.

    python scripts/make_sample_docx.py tests/fixtures/documents/sample.docx

This stands in until real documents are collected. It is deliberately written
the way people actually write these - a metadata table, bold pseudo-headings in
one variant, subjects both inline and in the heading - so the parser is exercised
against realistic mess rather than a clean schema.
"""

from __future__ import annotations

import sys
from pathlib import Path

import docx


def build(path: Path, use_real_headings: bool = True) -> Path:
    document = docx.Document()

    def heading(text: str, level: int = 2) -> None:
        if use_real_headings:
            document.add_heading(text, level=level)
        else:
            paragraph = document.add_paragraph()
            paragraph.add_run(text).bold = True

    heading("Senior Recruitment Consultant", 1)

    table = document.add_table(rows=0, cols=2)
    for label, value in (
        ("Location", "Manchester (hybrid)"),
        ("Salary", "£35,000 - £45,000 plus commission"),
        ("Employment Type", "Permanent"),
        ("Sector", "Technology"),
        ("Reference", "TR-4471"),
        ("Start Date", "Immediate"),
    ):
        row = table.add_row().cells
        row[0].text = label
        row[1].text = value

    heading("The Role")
    document.add_paragraph(
        "We are looking for a Senior Recruitment Consultant to join a growing "
        "technology desk. You will own the full 360 cycle, from business "
        "development through to offer management."
    )
    paragraph = document.add_paragraph()
    paragraph.add_run("You will be responsible for ").bold = False
    paragraph.add_run("building and owning your own desk").bold = True
    paragraph.add_run(", supported by a dedicated resourcing team.")

    for item in (
        "Manage a portfolio of existing clients and win new ones",
        "Source, screen and place technology candidates",
        "Hit and exceed quarterly billing targets",
    ):
        document.add_paragraph(item, style="List Bullet")

    heading("Email Sequence", 1)

    heading("Email 1 - Introduction")
    document.add_paragraph("Subject: Senior Recruitment Consultant - Manchester")
    document.add_paragraph("Hi {{first_name}},")
    document.add_paragraph(
        "I am working with a technology recruitment business in Manchester that "
        "is hiring a Senior Consultant. Given your background, I thought it was "
        "worth a conversation."
    )
    paragraph = document.add_paragraph()
    paragraph.add_run("The package is £35,000 to £45,000 plus ").bold = False
    paragraph.add_run("uncapped commission").bold = True
    paragraph.add_run(".")
    document.add_paragraph("Would you be open to a short call this week?")

    heading("Email 2 - Follow up (send after 3 days)")
    document.add_paragraph("Hi {{first_name}},")
    document.add_paragraph(
        "Following up on my note about the Senior Consultant role in Manchester. "
        "The team is interviewing this month."
    )
    for item in (
        "Hybrid working, two days in the office",
        "Dedicated resourcing support",
        "Clear route to team lead within 18 months",
    ):
        document.add_paragraph(item, style="List Bullet")
    document.add_paragraph("Happy to send the full spec across if useful.")

    heading("Email 3 - Final (day 7)")
    document.add_paragraph("Subject: Closing the loop")
    document.add_paragraph("Hi {{first_name}},")
    document.add_paragraph(
        "I have not heard back, so I will assume the timing is not right. If "
        "that changes, my details are below."
    )

    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(path))
    return path


if __name__ == "__main__":
    target = Path(sys.argv[1] if len(sys.argv) > 1 else "tests/fixtures/documents/sample.docx")
    real_headings = "--no-headings" not in sys.argv
    written = build(target, use_real_headings=real_headings)
    print(f"wrote {written}")
